import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import re
import os

class DocumentParser:
    """
    多模态文档解析器
    负责 PDF 与 Word 文档的文本提取、结构化转换及页面渲染。
    """
    def __init__(self):
        # 匹配标准分级标题正则表达式（如：1.1, 3.1.2）
        self.heading_re = re.compile(r'^(\d+(\.\d+)*)\s+(.*)')

    def parse_to_markdown(self, file_path):
        """
        根据扩展名选择解析方法并返回 Markdown 内容
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        return ""

    def _parse_pdf(self, path):
        """解析 PDF 文本与表格数据"""
        md_output = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                # 提取并格式化表格
                tables = page.extract_tables()
                for table in tables:
                    md_output.append(self._format_table(table))
                
                # 提取纯文本并识别层级标题
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        match = self.heading_re.match(line)
                        if match:
                            level = min(match.group(1).count('.') + 1, 6)
                            md_output.append(f"{'#' * level} {line}")
                        else:
                            md_output.append(line)
        return "\n\n".join(md_output)

    def _parse_docx(self, path):
        """解析 Word 文档段落与表格"""
        doc = Document(path)
        md_output = []
        for para in doc.paragraphs:
            # 识别 Word 内置的标题样式
            if para.style.name.startswith('Heading'):
                md_output.append(f"# {para.text}")
            else:
                md_output.append(para.text)
        # 遍历文档内的所有表格
        for table in doc.tables:
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            md_output.append(self._format_table(data))
        return "\n\n".join(md_output)

    def _format_table(self, data):
        """将二维列表数据转换为标准的 Markdown 表格格式"""
        if not data or not data[0]: return ""
        headers = [str(h) if h else "" for h in data[0]]
        md = f"| {' | '.join(headers)} |\n| {' | '.join(['---']*len(headers))} |\n"
        for row in data[1:]:
            row_str = [str(r) if r else "" for r in row]
            md += f"| {' | '.join(row_str)} |\n"
        return md

    def render_pdf_pages(self, pdf_path, output_folder):
        """将 PDF 每一页渲染为高清图片，用于差异比对的可视化展示"""
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        img_list = []
        with fitz.open(pdf_path) as doc:
            for i, page in enumerate(doc):
                # 渲染为 150 DPI 的图片以平衡清晰度与处理速度
                pix = page.get_pixmap(dpi=150)
                img_name = f"p_{i}.png"
                img_path = os.path.join(output_folder, img_name)
                pix.save(img_path)
                # 记录相对路径供 HTML 页面加载
                img_list.append(f"{os.path.basename(output_folder)}/{img_name}")
        return img_list